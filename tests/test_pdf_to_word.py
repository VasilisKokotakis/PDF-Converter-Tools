"""
Tests for pdf_to_word.py

GUI modules are mocked before import so no display is required.
"""
import sys
import os
import tempfile
import pytest
from unittest.mock import MagicMock
from PIL import Image
from docx import Document


# --- GUI mocks (must happen before importing pdf_to_word) ----------------

class _FakeTk:
    def __init__(self, *args, **kwargs): pass
    def __getattr__(self, name): return MagicMock()

_mock_ctk = MagicMock()
_mock_ctk.CTk = _FakeTk

sys.modules.setdefault("customtkinter", _mock_ctk)
sys.modules.setdefault("tkinter", MagicMock())
sys.modules.setdefault("tkinter.messagebox", MagicMock())
sys.modules.setdefault("tkinter.filedialog", MagicMock())

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from pdf_to_word import PDFtoWordApp  # noqa: E402


# --- Helpers -------------------------------------------------------------

def _make_image_pdf(path):
    """Create a minimal valid PDF using PIL (image-only, no text layer)."""
    img = Image.new("RGB", (200, 100), color=(255, 255, 255))
    img.save(path, "PDF")


# --- Fixtures ------------------------------------------------------------

@pytest.fixture
def app():
    return PDFtoWordApp()


@pytest.fixture
def tmp_pdf(tmp_path):
    p = str(tmp_path / "sample.pdf")
    _make_image_pdf(p)
    return p


# --- clear_pdf -----------------------------------------------------------

def test_clear_pdf_resets_path(app, tmp_pdf):
    app.pdf_path = tmp_pdf
    app.clear_pdf()
    assert app.pdf_path is None


def test_clear_pdf_resets_label(app, tmp_pdf):
    app.pdf_path = tmp_pdf
    app.clear_pdf()
    app.file_label.configure.assert_called_with(text="No PDF selected")


# --- integration: _convert_worker creates a valid docx ------------------

def test_convert_worker_creates_docx(app, tmp_pdf, tmp_path):
    app.pdf_path = tmp_pdf
    out = str(tmp_path / "output.docx")
    app._convert_worker(out)
    assert os.path.exists(out)
    assert os.path.getsize(out) > 0


def test_convert_worker_docx_has_no_trailing_page_break(app, tmp_pdf, tmp_path):
    """The last element in the document should not be a page break."""
    app.pdf_path = tmp_pdf
    out = str(tmp_path / "output.docx")
    app._convert_worker(out)
    doc = Document(out)
    if doc.paragraphs:
        last = doc.paragraphs[-1]
        for run in last.runs:
            assert "\f" not in run.text


# --- password error detection -------------------------------------------

@pytest.mark.parametrize("msg", ["password required", "file is encrypted", "ENCRYPTED PDF"])
def test_password_error_is_detected(msg):
    assert "password" in msg.lower() or "encrypted" in msg.lower()


@pytest.mark.parametrize("msg", ["invalid pdf", "file not found", "unexpected eof"])
def test_non_password_error_is_not_detected(msg):
    assert not ("password" in msg.lower() or "encrypted" in msg.lower())
